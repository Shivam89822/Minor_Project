import os
import re

import fitz
from docx import Document


def _split_text_blocks(text):
    if not text:
        return []

    normalized_text = text.replace("\r\n", "\n")
    raw_blocks = re.split(r"\n\s*\n+", normalized_text)
    blocks = []

    for raw_block in raw_blocks:
        cleaned_block = " ".join(line.strip() for line in raw_block.splitlines() if line.strip()).strip()
        if cleaned_block:
            blocks.append(cleaned_block)

    return blocks


def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)

    for page in doc:
        text += page.get_text() + "\n"

    doc.close()
    return text.strip()


def extract_pages_from_pdf(file_path):
    pages = []
    doc = fitz.open(file_path)

    for page_index, page in enumerate(doc, start=1):
        pages.append(
            {
                "page": page_index,
                "text": page.get_text().strip(),
            }
        )

    doc.close()
    return pages


def extract_paragraphs_from_pdf(file_path):
    paragraphs = []
    doc = fitz.open(file_path)
    paragraph_index = 0

    for page_index, page in enumerate(doc, start=1):
        blocks = sorted(page.get_text("blocks"), key=lambda block: (block[1], block[0]))

        for block in blocks:
            if len(block) < 5:
                continue

            for block_text in _split_text_blocks(block[4]):
                paragraph_index += 1
                paragraphs.append(
                    {
                        "page": page_index,
                        "section": None,
                        "paragraph_index": paragraph_index,
                        "text": block_text,
                    }
                )

    doc.close()
    return paragraphs


def extract_text_from_docx(file_path):
    text = ""
    doc = Document(file_path)

    for para in doc.paragraphs:
        text += para.text + "\n"

    return text.strip()


def extract_sections_from_docx(file_path):
    sections = []
    doc = Document(file_path)
    current_section = "Document"
    paragraph_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style and para.style.name else ""
        if "heading" in style_name:
            current_section = text
            continue

        paragraph_index += 1
        sections.append(
            {
                "page": None,
                "section": current_section,
                "paragraph_index": paragraph_index,
                "text": text,
            }
        )

    if not sections:
        sections.append(
            {
                "page": None,
                "section": "Document",
                "paragraph_index": 0,
                "text": "",
            }
        )

    return sections


def extract_document_units(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError("File does not exist")

    if file_path.endswith(".pdf"):
        paragraphs = extract_paragraphs_from_pdf(file_path)
        if paragraphs:
            return paragraphs

        pages = extract_pages_from_pdf(file_path)
        return [
            {
                "page": page["page"],
                "section": None,
                "paragraph_index": index,
                "text": page["text"],
            }
            for index, page in enumerate(pages, start=1)
        ]

    if file_path.endswith(".docx"):
        return extract_sections_from_docx(file_path)

    raise ValueError("Unsupported file format")


def extract_text(file_path):
    units = extract_document_units(file_path)
    return "\n".join(unit["text"] for unit in units if unit["text"]).strip()
